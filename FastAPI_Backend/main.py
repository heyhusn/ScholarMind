import io
import json
import os
import re
import urllib.parse
import urllib.request
import uuid
from typing import List, Optional

from fastapi import FastAPI, File, HTTPException, UploadFile, Form
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware

from ai import (
    analyze_paper,
    generate_chat_response,
    generate_flashcards,
    generate_podcast_script,
    generate_summary,
    simplify_text,
    generate_quiz_questions,
    generate_paper_insights,
    generate_peer_review,
    generate_paper_references,
)
from models import (
    AIRequest,
    AIResponse,
    ChatRequest,
    DocumentRequest,
    FlashcardItem,
    FlashcardResponse,
    PaperAnalysis,
    PaperSection,
    SimplifyRequest,
    TextResponse,
    QuizQuestion,
    QuizResponse,
    OpenAlexPaper,
    OpenAlexSearchResponse,
    PaperInsightsRequest,
    PaperInsightsResponse,
    PeerReviewResponse,
    ReferencesResponse,
    ExportReferencesRequest,
)
from pdf import extract_text_from_pdf, extract_text_from_docx

app = FastAPI(title="Scholar Mind Backend", description="Backend for Scholar Mind AI features")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

DOCUMENTS = {}
ANALYSES = {}


@app.get("/")
def read_root():
    return {"message": "Welcome to Scholar Mind AI Backend"}


def _strip_code_fences(raw_json: str) -> str:
    if raw_json.startswith("```json"):
        raw_json = raw_json[7:]
    if raw_json.startswith("```"):
        raw_json = raw_json[3:]
    if raw_json.endswith("```"):
        raw_json = raw_json[:-3]
    return raw_json.strip()


def _repair_json(raw_json: str) -> str:
    """Repair trailing commas and basic bracket closures in AI-generated JSON."""
    cleaned = _strip_code_fences(raw_json).strip()
    # Remove trailing commas before matching brace or bracket
    cleaned = re.sub(r',\s*([\]}])', r'\1', cleaned)
    # Basic unescaped control character replacement
    cleaned = re.sub(r'\t', ' ', cleaned)
    return cleaned


def _slugify_section_key(title: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", (title or "").strip().lower()).strip("_")
    return slug or "section"


def _build_sections(data: dict) -> List[PaperSection]:
    sections = []
    raw_sections = data.get("sections")
    if isinstance(raw_sections, list):
        for item in raw_sections:
            if not isinstance(item, dict):
                continue
            title = str(item.get("title", "")).strip()
            content = str(item.get("content", "")).strip()
            if not title or not content:
                continue
            key = str(item.get("key", "")).strip() or _slugify_section_key(title)
            sections.append(PaperSection(key=key, title=title, content=content))

    if sections:
        return sections

    legacy_fields = [
        ("abstract", "Abstract"),
        ("methodology", "Methodology"),
        ("results", "Results"),
        ("conclusion", "Conclusion"),
    ]
    for field_key, title in legacy_fields:
        content = data.get(field_key)
        if isinstance(content, str) and content.strip():
            sections.append(PaperSection(key=field_key, title=title, content=content.strip()))
    return sections


def _find_section_content(sections: List[PaperSection], preferred_key: str) -> Optional[str]:
    for section in sections:
        if section.key == preferred_key:
            return section.content
    return None


def first_nonempty_line(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped:
            return stripped
    return "Untitled Paper"


def build_fallback_analysis(doc_id: str, filename: str, extracted_text: str, error_message: str | None = None) -> PaperAnalysis:
    title = first_nonempty_line(extracted_text)
    year_match = re.search(r"\b(19|20)\d{2}\b", extracted_text)
    year = year_match.group(0) if year_match else ""

    abstract_text = None
    lowered = extracted_text.lower()
    abstract_idx = lowered.find("abstract")
    if abstract_idx != -1:
        abstract_text = extracted_text[abstract_idx:abstract_idx + 1200].strip()

    short_excerpt = " ".join(extracted_text.split())
    short_excerpt = short_excerpt[:500].strip()
    if not short_excerpt:
        short_excerpt = "Text was extracted, but a summary could not be generated."

    overview_body = "This analysis was generated using a fallback path because the AI response could not be parsed."
    if error_message:
        overview_body += f" Reason: {error_message[:180]}"

    sections = [
        PaperSection(key="abstract", title="Abstract", content=abstract_text or "No abstract content was extracted."),
        PaperSection(key="results", title="Results & Excerpt", content=short_excerpt)
    ]

    return PaperAnalysis(
        doc_id=doc_id,
        title=title if title else filename,
        authors="Unknown Authors",
        year=year,
        venue="Unknown",
        field="Research",
        citation_count=0,
        citation_impact="Low",
        citation_score=0,
        ai_overview_title="Fallback analysis generated",
        ai_overview_body=overview_body,
        sections=sections,
        abstract=abstract_text,
        methodology=None,
        results=short_excerpt,
        conclusion=None,
        citations_list=None,
    )


@app.post("/api/pdf/extract")
async def process_pdf(file: UploadFile = File(...)):
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")

    extracted_text = await extract_text_from_pdf(file)
    if extracted_text.startswith("Error"):
        raise HTTPException(status_code=500, detail=extracted_text)

    doc_id = str(uuid.uuid4())
    DOCUMENTS[doc_id] = extracted_text
    return {"doc_id": doc_id, "filename": file.filename, "message": "Document processed and stored."}


@app.post("/api/pdf/analyze", response_model=PaperAnalysis)
async def analyze_pdf(file: UploadFile = File(...), user_id: str = Form("")):
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")

    extracted_text = await extract_text_from_pdf(file)
    if extracted_text.startswith("Error"):
        raise HTTPException(status_code=500, detail=extracted_text)

    doc_id = str(uuid.uuid4())
    DOCUMENTS[doc_id] = extracted_text

    raw_json = analyze_paper(extracted_text)
    if raw_json.startswith("Error:"):
        raise HTTPException(status_code=500, detail=raw_json)

    analysis_error = None
    try:
        repaired = _repair_json(raw_json)
        data = json.loads(repaired)
        sections = _build_sections(data)
        analysis = PaperAnalysis(
            doc_id=doc_id,
            title=data.get("title", file.filename),
            authors=data.get("authors", "Unknown Authors"),
            year=data.get("year", ""),
            venue=data.get("venue", ""),
            field=data.get("field", "Research"),
            citation_count=int(data.get("citation_count", 0)),
            citation_impact=data.get("citation_impact", "Unknown"),
            citation_score=int(data.get("citation_score", 0)),
            ai_overview_title=data.get("ai_overview_title", ""),
            ai_overview_body=data.get("ai_overview_body", ""),
            sections=sections,
            abstract=_find_section_content(sections, "abstract"),
            methodology=_find_section_content(sections, "methodology"),
            results=_find_section_content(sections, "results"),
            conclusion=_find_section_content(sections, "conclusion"),
            citations_list=data.get("citations_list"),
        )
    except Exception as e:
        analysis_error = str(e)
        print(f"AI analysis parsing failed for doc {doc_id}: {analysis_error}. Running fallback.")
        analysis = build_fallback_analysis(doc_id, file.filename, extracted_text, analysis_error)

    ANALYSES[doc_id] = analysis.dict()
    return analysis


@app.get("/api/pdf/analysis/{doc_id}", response_model=PaperAnalysis)
def get_analysis(doc_id: str, user_id: Optional[str] = None):
    if doc_id in ANALYSES:
        return ANALYSES[doc_id]

    firestore_analysis = _paper_analysis_from_firestore(doc_id, user_id)
    if firestore_analysis is not None:
        return firestore_analysis

    raise HTTPException(status_code=404, detail="Analysis not found for this doc_id.")


def get_firebase_project_id() -> str:
    paths = [
        "c:/Users/SNAKE/AndroidStudioProjects/ScholarApp/app/google-services.json",
        "C:/Users/SNAKE/AndroidStudioProjects/ScholarApp/app/google-services.json",
        "C:/Users/SNAKE/Desktop/mob app/scholarai/google-services.json",
        "../google-services.json",
        "google-services.json"
    ]
    for p in paths:
        if os.path.exists(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    p_id = data.get("project_info", {}).get("project_id")
                    if p_id:
                        return p_id
            except Exception:
                pass
    return "scholarapp-43ed5"


def get_document_context(doc_id: str, user_id: Optional[str] = None) -> str:
    if doc_id in DOCUMENTS:
        return DOCUMENTS[doc_id]

    if user_id:
        project_id = get_firebase_project_id()
        url = f"https://firestore.googleapis.com/v1/projects/{project_id}/databases/(default)/documents/users/{user_id}/papers/{doc_id}"
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=5) as response:
                data = json.loads(response.read().decode())

            fields = data.get("fields", {})
            sections_val = fields.get("sections", {})

            text_parts = []
            array_val = sections_val.get("arrayValue", {})
            values = array_val.get("values", [])

            for val in values:
                map_val = val.get("mapValue", {})
                m_fields = map_val.get("fields", {})
                title = m_fields.get("title", {}).get("stringValue", "")
                content = m_fields.get("content", {}).get("stringValue", "")
                if title:
                    text_parts.append(f"Section: {title}")
                if content:
                    text_parts.append(content)

            if not text_parts:
                for field_key in ["abstract", "methodology", "results", "conclusion"]:
                    field_val = fields.get(field_key, {}).get("stringValue", "")
                    if field_val:
                        text_parts.append(f"Section: {field_key.capitalize()}")
                        text_parts.append(field_val)

            if text_parts:
                context = "\n\n".join(text_parts)
                DOCUMENTS[doc_id] = context
                return context
        except Exception as e:
            print(f"Firestore fallback failed for doc {doc_id}, user {user_id}: {e}")

    raise HTTPException(status_code=404, detail="Document not found. Cache is empty and Firestore fetch failed.")


def _firestore_string_value(fields: dict, key: str, default: str = "") -> str:
    value = fields.get(key, {})
    if not isinstance(value, dict):
        return default

    if "stringValue" in value and value["stringValue"] is not None:
        return str(value["stringValue"])
    if "integerValue" in value and value["integerValue"] is not None:
        return str(value["integerValue"])
    if "doubleValue" in value and value["doubleValue"] is not None:
        return str(value["doubleValue"])
    if "booleanValue" in value and value["booleanValue"] is not None:
        return "true" if value["booleanValue"] else "false"
    return default


def _firestore_int_value(fields: dict, key: str, default: int = 0) -> int:
    value = fields.get(key, {})
    if not isinstance(value, dict):
        return default

    for field_name in ("integerValue", "doubleValue", "stringValue"):
        raw = value.get(field_name)
        if raw is None or raw == "":
            continue
        try:
            return int(float(raw))
        except Exception:
            continue
    return default


def _paper_analysis_from_firestore(doc_id: str, user_id: Optional[str]) -> Optional[PaperAnalysis]:
    if not user_id:
        return None

    project_id = get_firebase_project_id()
    url = f"https://firestore.googleapis.com/v1/projects/{project_id}/databases/(default)/documents/users/{user_id}/papers/{doc_id}"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=5) as response:
            data = json.loads(response.read().decode())

        fields = data.get("fields", {})
        section_values = []
        sections_field = fields.get("sections", {})
        array_values = sections_field.get("arrayValue", {}).get("values", [])

        for item in array_values:
            map_fields = item.get("mapValue", {}).get("fields", {})
            title = _firestore_string_value(map_fields, "title", "")
            content = _firestore_string_value(map_fields, "content", "")
            key = _firestore_string_value(map_fields, "key", "") or _slugify_section_key(title)
            if title and content:
                section_values.append({"key": key, "title": title, "content": content})

        if not section_values:
            legacy_sections = []
            for field_key, title in [
                ("abstract", "Abstract"),
                ("methodology", "Methodology"),
                ("results", "Results"),
                ("conclusion", "Conclusion"),
            ]:
                content = _firestore_string_value(fields, field_key, "")
                if content.strip():
                    legacy_sections.append({"key": field_key, "title": title, "content": content.strip()})
            section_values = legacy_sections

        analysis = PaperAnalysis(
            doc_id=_firestore_string_value(fields, "docId", doc_id) or doc_id,
            title=_firestore_string_value(fields, "title", "Untitled"),
            authors=_firestore_string_value(fields, "authors", _firestore_string_value(fields, "author", "Unknown Authors")),
            year=_firestore_string_value(fields, "year", ""),
            venue=_firestore_string_value(fields, "venue", ""),
            field=_firestore_string_value(fields, "field", "Research"),
            citation_count=_firestore_int_value(fields, "citationCount", 0),
            citation_impact=_firestore_string_value(fields, "citationImpact", "Unknown"),
            citation_score=_firestore_int_value(fields, "citationScore", 0),
            ai_overview_title=_firestore_string_value(fields, "aiOverviewTitle", ""),
            ai_overview_body=_firestore_string_value(fields, "aiOverviewBody", ""),
            sections=[PaperSection(**section) for section in section_values],
            abstract=_firestore_string_value(fields, "abstract", None),
            methodology=_firestore_string_value(fields, "methodology", None),
            results=_firestore_string_value(fields, "results", None),
            conclusion=_firestore_string_value(fields, "conclusion", None),
            citations_list=_firestore_string_value(fields, "citationsList", None),
        )

        ANALYSES[doc_id] = analysis.dict()
        return analysis
    except Exception as e:
        print(f"Firestore analysis fallback failed for doc {doc_id}, user {user_id}: {e}")
        return None


@app.post("/api/chat/beginner", response_model=TextResponse)
def chat_beginner(request: ChatRequest):
    context = get_document_context(request.doc_id, request.user_id)
    return TextResponse(text=generate_chat_response(context, request.message, "beginner"))


@app.post("/api/chat/technical", response_model=TextResponse)
def chat_technical(request: ChatRequest):
    context = get_document_context(request.doc_id, request.user_id)
    return TextResponse(text=generate_chat_response(context, request.message, "technical"))


@app.post("/api/chat/freeform", response_model=TextResponse)
def chat_freeform(request: ChatRequest):
    context = get_document_context(request.doc_id, request.user_id)
    return TextResponse(text=generate_chat_response(context, request.message, "freeform"))


@app.post("/api/flashcards/generate", response_model=FlashcardResponse)
def flashcards_generate(request: DocumentRequest):
    context = get_document_context(request.doc_id, request.user_id)
    json_response = generate_flashcards(context)
    if json_response.startswith("Error"):
        raise HTTPException(status_code=500, detail=json_response)
    try:
        repaired = _repair_json(json_response)
        flashcards_data = json.loads(repaired)
        flashcards = [FlashcardItem(**item) for item in flashcards_data]
        return FlashcardResponse(flashcards=flashcards)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing flashcards: {e}")


@app.post("/api/podcast/generate", response_model=TextResponse)
def podcast_generate(request: DocumentRequest):
    context = get_document_context(request.doc_id, request.user_id)
    return TextResponse(text=generate_podcast_script(context))


@app.post("/api/summary/generate", response_model=TextResponse)
def summary_generate(request: DocumentRequest):
    context = get_document_context(request.doc_id, request.user_id)
    return TextResponse(text=generate_summary(context))


@app.post("/api/text/simplify", response_model=TextResponse)
def text_simplify(request: SimplifyRequest):
    return TextResponse(text=simplify_text(request.text))


@app.post("/api/quiz/generate", response_model=QuizResponse)
def quiz_generate(request: DocumentRequest):
    context = get_document_context(request.doc_id, request.user_id)
    json_response = generate_quiz_questions(context)
    if json_response.startswith("Error"):
        raise HTTPException(status_code=500, detail=json_response)
    
    try:
        repaired = _repair_json(json_response)
        quiz_data = json.loads(repaired)
        questions = [QuizQuestion(**item) for item in quiz_data]
        return QuizResponse(questions=questions)
    except Exception as e:
        print(f"Standard quiz parsing failed: {e}. Attempting regex fallback.")
        try:
            quiz_data = []
            # Find and isolate question blocks
            q_blocks = re.split(r'\{\s*"id"', json_response)
            for idx, block in enumerate(q_blocks[1:], start=1):
                q_id = f"q{idx}"
                id_match = re.search(r'"id"\s*:\s*"([^"]+)"', block)
                if id_match:
                    q_id = id_match.group(1)
                
                question = "Comprehension Question"
                q_match = re.search(r'"question"\s*:\s*"([^"]+)"', block)
                if q_match:
                    question = q_match.group(1)
                
                # Options
                options = []
                opt_match = re.search(r'"options"\s*:\s*\[(.*?)\]', block, re.DOTALL)
                if opt_match:
                    options = re.findall(r'"([^"]+)"', opt_match.group(1))
                if len(options) < 4:
                    options += ["Option A", "Option B", "Option C", "Option D"][len(options):]
                
                correct_idx = 0
                idx_match = re.search(r'"correct_answer_index"\s*:\s*(\d+)', block)
                if idx_match:
                    correct_idx = int(idx_match.group(1))
                if correct_idx >= len(options):
                    correct_idx = 0
                    
                explanation = "Refer to the text context."
                exp_match = re.search(r'"explanation"\s*:\s*"([^"]+)"', block)
                if exp_match:
                    explanation = exp_match.group(1)
                    
                quiz_data.append(QuizQuestion(
                    id=q_id,
                    question=question,
                    options=options[:4],
                    correct_answer_index=correct_idx,
                    explanation=explanation
                ))
            if quiz_data:
                return QuizResponse(questions=quiz_data)
        except Exception as ex:
            print(f"Quiz regex fallback also failed: {ex}")
        
        raise HTTPException(status_code=500, detail=f"Error processing quiz JSON: {e}")


def _reconstruct_abstract(inverted_index: Optional[dict]) -> Optional[str]:
    if not inverted_index:
        return None
    try:
        max_idx = -1
        for word, indices in inverted_index.items():
            for idx in indices:
                if idx > max_idx:
                    max_idx = idx
        if max_idx == -1:
            return ""
        words = [""] * (max_idx + 1)
        for word, indices in inverted_index.items():
            for idx in indices:
                words[idx] = word
        return " ".join(words)
    except Exception:
        return None


def _parse_openalex_work(work: dict) -> OpenAlexPaper:
    raw_id = work.get("id", "")
    paper_id = raw_id.split("/")[-1] if "/" in raw_id else raw_id
    
    authors_list = []
    authorships = work.get("authorships", [])
    for auth in authorships:
        author_name = auth.get("author", {}).get("display_name")
        if author_name:
            authors_list.append(author_name)
    authors_str = ", ".join(authors_list) if authors_list else "Unknown Authors"
    
    primary_location = work.get("primary_location", {}) or {}
    source = primary_location.get("source", {}) or {}
    venue = source.get("display_name") or "Unknown Venue"
    
    primary_topic = work.get("primary_topic", {}) or {}
    topic_name = primary_topic.get("display_name") or "Research"
    
    open_access = work.get("open_access", {}) or {}
    is_oa = bool(open_access.get("is_oa", False))
    oa_pdf = open_access.get("oa_url") or None
    
    doi = work.get("doi")
    abstract = _reconstruct_abstract(work.get("abstract_inverted_index"))
    
    # Extra fields:
    # 1. Keywords
    keywords_list = [k.get("display_name") for k in work.get("keywords", []) if k.get("display_name")]
    keywords_str = ", ".join(keywords_list[:5]) if keywords_list else "None"
    
    # 2. Publisher
    publisher = source.get("publisher") or "Unknown Publisher"
    
    # 3. Funders & Awards
    grants = work.get("grants") or []
    funders = ", ".join(list(set([g.get("funder_display_name") for g in grants if g.get("funder_display_name")]))) or "None"
    awards = ", ".join(list(set([g.get("award_id") for g in grants if g.get("award_id")]))) or "None"
    
    # 4. Domain, Field, Subfield
    domain = primary_topic.get("domain", {}).get("display_name") or "Unknown Domain"
    field_name = primary_topic.get("field", {}).get("display_name") or "Unknown Field"
    subfield = primary_topic.get("subfield", {}).get("display_name") or "Unknown Subfield"
    
    # 5. SDGs
    sdgs_list = [s.get("display_name") for s in work.get("sdgs", []) if s.get("display_name")]
    sdgs_str = ", ".join(sdgs_list) if sdgs_list else "None"
    
    # 6. Countries & Continents
    countries_set = set()
    continents_set = set()
    for auth in authorships:
        countries = auth.get("countries") or []
        for c in countries:
            countries_set.add(c)
        institutions = auth.get("institutions") or []
        for inst in institutions:
            c_code = inst.get("country_code")
            if c_code:
                countries_set.add(c_code)
            cont = inst.get("continent")
            if cont:
                continents_set.add(cont)
                
    countries_str = ", ".join(sorted(list(countries_set))) if countries_set else "Unknown"
    continents_str = ", ".join(sorted(list(continents_set))) if continents_set else "Unknown"
    
    # 7. Language
    language = work.get("language") or "Unknown"
    
    return OpenAlexPaper(
        id=paper_id,
        title=work.get("title") or "Untitled",
        authors=authors_str,
        year=str(work.get("publication_year") or ""),
        venue=venue,
        citation_count=work.get("cited_by_count") or 0,
        doi=doi,
        is_open_access=is_oa,
        open_access_pdf=oa_pdf,
        primary_topic=topic_name,
        abstract=abstract,
        keywords=keywords_str,
        publisher=publisher,
        funders=funders,
        awards=awards,
        domain=domain,
        field_name=field_name,
        subfield=subfield,
        sdgs=sdgs_str,
        countries=countries_str,
        continents=continents_str,
        language=language
    )


@app.get("/api/papers/search", response_model=OpenAlexSearchResponse)
def search_papers(query: str, page: int = 1, filter: Optional[str] = None):
    base_url = "https://api.openalex.org/works"
    params = {
        "search": query,
        "page": str(page),
        "per_page": "20"
    }
    if filter:
        params["filter"] = filter
    
    key = os.getenv("OPENALEX_API_KEY") or os.getenv("OPENSALEX_API_KEY")
    if key:
        params["api_key"] = key
        
    query_str = urllib.parse.urlencode(params)
    url = f"{base_url}?{query_str}"
    
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "ScholarMind/1.0 (mailto:scholarapp@example.com)"})
        with urllib.request.urlopen(req, timeout=10) as response:
            data = json.loads(response.read().decode("utf-8"))
            
        meta = data.get("meta", {})
        results = data.get("results", [])
        
        parsed_results = [_parse_openalex_work(w) for w in results]
        
        return OpenAlexSearchResponse(
            count=meta.get("count") or 0,
            page=meta.get("page") or page,
            per_page=meta.get("per_page") or 20,
            results=parsed_results
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAlex request failed: {e}")


@app.get("/api/papers/trending", response_model=OpenAlexSearchResponse)
def get_trending_papers():
    base_url = "https://api.openalex.org/works"
    params = {
        "filter": "publication_year:2024|2025,cited_by_count:>5",
        "sort": "cited_by_count:desc",
        "page": "1",
        "per_page": "15"
    }
    key = os.getenv("OPENALEX_API_KEY") or os.getenv("OPENSALEX_API_KEY")
    if key:
        params["api_key"] = key
        
    query_str = urllib.parse.urlencode(params)
    url = f"{base_url}?{query_str}"
    
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "ScholarMind/1.0 (mailto:scholarapp@example.com)"})
        with urllib.request.urlopen(req, timeout=10) as response:
            data = json.loads(response.read().decode("utf-8"))
            
        meta = data.get("meta", {})
        results = data.get("results", [])
        
        parsed_results = [_parse_openalex_work(w) for w in results]
        
        return OpenAlexSearchResponse(
            count=meta.get("count") or 0,
            page=1,
            per_page=15,
            results=parsed_results
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAlex trending request failed: {e}")


@app.get("/api/papers/{paper_id}", response_model=OpenAlexPaper)
def get_paper_details(paper_id: str):
    url = f"https://api.openalex.org/works/{paper_id}"
    key = os.getenv("OPENALEX_API_KEY") or os.getenv("OPENSALEX_API_KEY")
    if key:
        url += f"?api_key={key}"
        
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "ScholarMind/1.0 (mailto:scholarapp@example.com)"})
        with urllib.request.urlopen(req, timeout=10) as response:
            work = json.loads(response.read().decode("utf-8"))
            
        return _parse_openalex_work(work)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAlex get paper failed: {e}")


@app.post("/api/papers/insights", response_model=PaperInsightsResponse)
def get_paper_insights(request: PaperInsightsRequest):
    raw_json = generate_paper_insights(request.title, request.abstract, request.authors)
    if raw_json.startswith("Error:"):
        raise HTTPException(status_code=500, detail=raw_json)
        
    try:
        repaired = _repair_json(raw_json)
        data = json.loads(repaired)
        return PaperInsightsResponse(
            summary=data.get("summary") or "No summary available.",
            key_findings=data.get("key_findings") or [],
            applications=data.get("applications") or [],
            limitations=data.get("limitations") or []
        )
    except Exception as e:
        print(f"Standard insights JSON parsing failed: {e}. Running regex fallback.")
        try:
            summary_match = re.search(r'"summary"\s*:\s*"([^"]+)"', raw_json)
            summary = summary_match.group(1) if summary_match else "No summary available."
            
            data = {}
            for key in ["key_findings", "applications", "limitations"]:
                list_match = re.search(rf'"{key}"\s*:\s*\[(.*?)\]', raw_json, re.DOTALL)
                items = []
                if list_match:
                    items = re.findall(r'"([^"]+)"', list_match.group(1))
                else:
                    block_match = re.search(rf'"{key}"\s*:\s*\[(.*?)(?:"\w+"\s*:|$)', raw_json, re.DOTALL)
                    if block_match:
                        items = re.findall(r'"([^"]+)"', block_match.group(1))
                data[key] = items
                
            return PaperInsightsResponse(
                summary=summary,
                key_findings=data.get("key_findings") or [],
                applications=data.get("applications") or [],
                limitations=data.get("limitations") or []
            )
        except Exception as ex:
            print(f"Insights regex fallback failed: {ex}")
            
        raise HTTPException(status_code=500, detail=f"Failed to generate valid insights JSON: {e}")


@app.post("/api/peer-review/analyze", response_model=PeerReviewResponse)
async def analyze_peer_review_endpoint(file: UploadFile = File(...)):
    filename = file.filename.lower()
    
    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
        
    try:
        if filename.endswith(".pdf"):
            text = await extract_text_from_pdf(file)
        else:
            file_bytes = await file.read()
            text = extract_text_from_docx(file_bytes)
            
        if text.startswith("Error"):
            raise HTTPException(status_code=500, detail=text)
            
        raw_json = generate_peer_review(text)
        if raw_json.startswith("Error:"):
            raise HTTPException(status_code=500, detail=raw_json)
            
        try:
            repaired = _repair_json(raw_json)
            data = json.loads(repaired)
            return PeerReviewResponse(
                limitations=data.get("limitations") or [],
                technical_flaws=data.get("technical_flaws") or [],
                questions=data.get("questions") or []
            )
        except Exception as e:
            print(f"Standard peer review parsing failed: {e}. Running regex fallback.")
            try:
                data = {}
                for key in ["limitations", "technical_flaws", "questions"]:
                    list_match = re.search(rf'"{key}"\s*:\s*\[(.*?)\]', raw_json, re.DOTALL)
                    items = []
                    if list_match:
                        items = re.findall(r'"([^"]+)"', list_match.group(1))
                    else:
                        block_match = re.search(rf'"{key}"\s*:\s*\[(.*?)(?:"\w+"\s*:|$)', raw_json, re.DOTALL)
                        if block_match:
                            items = re.findall(r'"([^"]+)"', block_match.group(1))
                    data[key] = items
                    
                return PeerReviewResponse(
                    limitations=data.get("limitations") or [],
                    technical_flaws=data.get("technical_flaws") or [],
                    questions=data.get("questions") or []
                )
            except Exception as ex:
                print(f"Peer review regex fallback failed: {ex}")
                
            raise HTTPException(status_code=500, detail=f"Failed to parse peer review JSON: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate peer review report: {e}")


@app.post("/api/references/generate", response_model=ReferencesResponse)
async def generate_references_endpoint(file: UploadFile = File(...), style: str = Form("APA")):
    filename = file.filename.lower()
    
    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
        
    try:
        if filename.endswith(".pdf"):
            text = await extract_text_from_pdf(file)
        else:
            file_bytes = await file.read()
            text = extract_text_from_docx(file_bytes)
            
        if text.startswith("Error"):
            raise HTTPException(status_code=500, detail=text)
            
        raw_json = generate_paper_references(text, style)
        if raw_json.startswith("Error:"):
            raise HTTPException(status_code=500, detail=raw_json)
            
        try:
            repaired = _repair_json(raw_json)
            data = json.loads(repaired)
        except Exception as e:
            print(f"Standard references parsing failed: {e}. Trying list parsing fallback.")
            repaired = re.sub(r',\s*([\]}])', r'\1', raw_json)
            try:
                data = json.loads(repaired)
            except Exception:
                # Regex list fallback
                data = re.findall(r'"([^"]+)"', raw_json)
                if not data:
                    raise e
        
        if not isinstance(data, list):
            if isinstance(data, dict) and "references" in data:
                data = data["references"]
            else:
                raise ValueError("AI response is not a JSON list of references")
                
        return ReferencesResponse(references=[str(r) for r in data])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate references: {e}")


@app.post("/api/references/export/docx")
def export_references_docx(request: ExportReferencesRequest):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library is not installed on the server.")
        
    try:
        doc = Document()
        
        # Page setup - 1 inch margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(f"References ({request.style} Style)")
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(14)
        title_run.bold = True
        title_p.paragraph_format.space_after = Pt(18)
        
        # References
        for ref in request.references:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(8)
            
            run = p.add_run(ref)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        headers = {
            "Content-Disposition": f'attachment; filename="references_{request.style.lower()}.docx"'
        }
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to export DOCX: {e}")


@app.post("/api/references/export/pdf")
def export_references_pdf(request: ExportReferencesRequest):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab library is not installed on the server.")
        
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontName='Times-Bold',
            fontSize=14,
            leading=16,
            alignment=TA_CENTER,
            spaceAfter=18
        )
        
        ref_style = ParagraphStyle(
            'RefItem',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=11,
            leading=14,
            alignment=TA_LEFT,
            leftIndent=36,
            firstLineIndent=-36,
            spaceAfter=8
        )
        
        story = []
        story.append(Paragraph(f"References ({request.style} Style)", title_style))
        
        for ref in request.references:
            escaped_ref = ref.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(escaped_ref, ref_style))
            
        doc.build(story)
        buffer.seek(0)
        
        headers = {
            "Content-Disposition": f'attachment; filename="references_{request.style.lower()}.pdf"'
        }
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers=headers
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to export PDF: {e}")
